import io
import json
import textwrap
import re
from datetime import datetime
from typing import Any, Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.core.exceptions import ValidationError
from app.core.templates import registry

SUPPORTED_EXPORT_FORMATS = {"pdf", "docx", "txt", "json"}

EXPORT_MIME_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "json": "application/json",
}

SECTION_LABELS = {
    "summary": "Perfil Profesional",
    "experience": "Experiencia",
    "education": "Educación",
    "skills": "Habilidades",
    "projects": "Proyectos",
    "certifications": "Certificaciones",
    "languages": "Idiomas",
    "interests": "Intereses",
    "tools": "Herramientas",
}


def _strip_font_name(value: Optional[str]) -> Optional[str]:
    if not value:
        return None
    return str(value).replace('"', '').strip()


def _clamp(value: float, minimum: float, maximum: float) -> float:
    return max(minimum, min(maximum, value))


def _normalize_hex(value: str) -> Optional[str]:
    if not value:
        return None
    match = re.match(r"^#([0-9a-fA-F]{3}|[0-9a-fA-F]{6})$", value)
    if not match:
        return None
    hex_value = value.lower()
    if len(hex_value) == 4:
        hex_value = "#" + "".join([c * 2 for c in hex_value[1:]])
    return hex_value


def _resolve_hex_color(value: Optional[str], fallback: str) -> str:
    if not value:
        return fallback
    raw = str(value).strip()
    if raw in OKLCH_TO_HEX:
        return OKLCH_TO_HEX[raw]
    normalized = _normalize_hex(raw)
    if normalized:
        return normalized
    return fallback


def _get_section_label(cv_data: Dict[str, Any], section: str) -> str:
    config = cv_data.get("config") or {}
    sections_config = config.get("sections") or {}
    section_config = sections_config.get(section) or {}
    return section_config.get("title") or SECTION_LABELS.get(section, section.title())


def _get_layout_config(cv_data: Dict[str, Any]) -> Dict[str, float]:
    config = cv_data.get("config") or {}
    layout = config.get("layout") or {}
    section_gap = layout.get("sectionGap", DEFAULT_LAYOUT["sectionGap"])
    content_gap = layout.get("contentGap", DEFAULT_LAYOUT["contentGap"])
    font_scale = _clamp(float(layout.get("fontSize", DEFAULT_LAYOUT["fontSize"])), 0.7, 1.3)
    return {
        "sectionGap": float(section_gap),
        "contentGap": float(content_gap),
        "fontScale": font_scale,
    }


def _resolve_docx_fonts(cv_data: Dict[str, Any], template_id: str) -> Dict[str, str]:
    config = cv_data.get("config") or {}
    fonts = config.get("fonts") or {}
    
    template_config = registry.get_template(template_id)
    default_heading = template_config.styles.get("heading_font", "Helvetica-Bold") if template_config else "Helvetica-Bold"
    default_body = template_config.styles.get("body_font", "Helvetica") if template_config else "Helvetica"
    
    heading = _strip_font_name(fonts.get("heading")) or default_heading
    body = _strip_font_name(fonts.get("body")) or default_body
    return {"heading": heading, "body": body}


def _resolve_pdf_fonts(cv_data: Dict[str, Any], template_id: str) -> Dict[str, str]:
    fonts = _resolve_docx_fonts(cv_data, template_id)

    def _map_pdf_font(name: str, bold: bool) -> str:
        font_name = name.lower()
        if "mono" in font_name or "code" in font_name or "courier" in font_name:
            return "Courier-Bold" if bold else "Courier"
        if "playfair" in font_name or "times" in font_name or "serif" in font_name:
            return "Times-Bold" if bold else "Times-Roman"
        return "Helvetica-Bold" if bold else "Helvetica"

    return {
        "heading": _map_pdf_font(fonts["heading"], True),
        "body": _map_pdf_font(fonts["body"], False),
    }


def _resolve_colors(cv_data: Dict[str, Any], template_id: str) -> Dict[str, str]:
    config = cv_data.get("config") or {}
    colors_config = config.get("colors") or {}
    
    template_config = registry.get_template(template_id)
    accent_fallback = "#111827"
    
    if template_config and template_config.styles.get("accent"):
        raw_accent = template_config.styles["accent"]
        # Handle ReportLab color objects if present (legacy) or hex strings
        if hasattr(raw_accent, 'hexval'):
             accent_fallback = f"#{raw_accent.hexval()[2:]}"
        elif isinstance(raw_accent, str):
             accent_fallback = raw_accent
    
    accent = _resolve_hex_color(colors_config.get("accent") or colors_config.get("primary"), accent_fallback)
    primary = _resolve_hex_color(colors_config.get("primary"), accent)
    text = _resolve_hex_color(colors_config.get("text"), "#111827")
    return {
        "accent": accent,
        "primary": primary,
        "text": text,
    }


def _px_to_pt(px: float) -> float:
    return float(px) * 0.75


def _text_color_to_rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value.replace("#", "").upper())


def _apply_run_style(run, font_name: str, font_size: float, color_hex: Optional[str] = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color_hex:
        run.font.color.rgb = _text_color_to_rgb(color_hex)


def _apply_paragraph_spacing(paragraph, after_pt: float) -> None:
    paragraph.paragraph_format.space_after = Pt(after_pt)


def _section_gap_lines(section_gap_px: float) -> int:
    if section_gap_px >= 56:
        return 3
    if section_gap_px >= 40:
        return 2
    return 1


def _content_gap_lines(content_gap_px: float) -> int:
    return 1 if content_gap_px >= 32 else 0


def build_export_payload(
    cv_data: Dict[str, Any],
    template_id: str,
    export_format: str,
) -> Tuple[bytes, str, str]:
    """Genera un archivo descargable con formato y plantilla específicos."""
    template_config = registry.get_template(template_id)
    if not template_config:
        raise ValidationError(f"Plantilla no soportada: {template_id}")

    if export_format not in SUPPORTED_EXPORT_FORMATS:
        raise ValidationError(f"Formato no soportado: {export_format}")

    file_name = _build_filename(cv_data, template_id, export_format)
    if export_format == "json":
        content = _build_json_export(cv_data, template_id)
    elif export_format == "txt":
        content = _build_txt_export(cv_data, template_id)
    elif export_format == "docx":
        content = _build_docx_export(cv_data, template_id)
    elif export_format == "pdf":
        content = _build_pdf_export(cv_data, template_id)
    else:
        raise ValidationError(f"Formato no soportado: {export_format}")

    return content, file_name, EXPORT_MIME_TYPES[export_format]


def _build_filename(cv_data: Dict[str, Any], template_id: str, export_format: str) -> str:
    personal_info = cv_data.get("personalInfo", {})
    full_name = personal_info.get("fullName") or "cv"
    safe_name = "_".join(full_name.strip().split())
    
    template_config = registry.get_template(template_id)
    template_label = template_config.name if template_config else template_id
    
    timestamp = datetime.utcnow().strftime("%Y%m%d")
    return f"cv_{safe_name}_{template_label}_{timestamp}.{export_format}"


def _is_section_visible(cv_data: Dict[str, Any], section: str) -> bool:
    config = cv_data.get("config") or {}
    sections_config = config.get("sections") or {}
    section_config = sections_config.get(section) or {}
    return section_config.get("visible", True)


def _build_json_export(cv_data: Dict[str, Any], template_id: str) -> bytes:
    payload = {
        "template": template_id,
        "generatedAt": datetime.utcnow().isoformat(),
        "data": cv_data,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def _build_txt_export(cv_data: Dict[str, Any], template_id: str) -> bytes:
    lines: List[str] = []
    personal_info = cv_data.get("personalInfo", {})
    full_name = personal_info.get("fullName") or "Sin nombre"
    role = personal_info.get("role")
    layout = _get_layout_config(cv_data)
    section_gap_lines = _section_gap_lines(layout["sectionGap"])
    content_gap_lines = _content_gap_lines(layout["contentGap"])

    lines.append(full_name)
    if role:
        lines.append(role)
    lines.append("=" * 60)

    contact_parts = [
        personal_info.get("email"),
        personal_info.get("phone"),
        personal_info.get("location"),
    ]
    contact = " | ".join([item for item in contact_parts if item])
    if contact:
        lines.append(contact)

    link_parts = []
    if personal_info.get("linkedin"):
        link_parts.append(f"LinkedIn: {personal_info.get('linkedin')}")
    if personal_info.get("github"):
        link_parts.append(f"GitHub: {personal_info.get('github')}")
    if personal_info.get("website"):
        link_parts.append(f"Web: {personal_info.get('website')}")
    if link_parts:
        lines.append(" | ".join(link_parts))

    lines.append("")

    for section in TEMPLATE_SECTION_ORDER[template_id]:
        if not _is_section_visible(cv_data, section):
            continue
        section_lines = _render_section_txt(cv_data, section, content_gap_lines)
        if section_lines:
            lines.extend(section_lines)
            lines.extend([""] * section_gap_lines)

    return "\n".join(lines).strip().encode("utf-8")


def _render_section_txt(cv_data: Dict[str, Any], section: str, content_gap_lines: int) -> List[str]:
    lines: List[str] = []
    label = _get_section_label(cv_data, section)

    if section == "summary":
        summary = cv_data.get("personalInfo", {}).get("summary")
        if summary:
            lines.append(label.upper())
            lines.append("-" * len(label))
            lines.append(summary)
    elif section == "experience":
        experiences = cv_data.get("experience", [])
        if experiences:
            lines.append(label.upper())
            lines.append("-" * len(label))
            for exp in experiences:
                title = exp.get("position") or "Sin título"
                company = exp.get("company")
                location = exp.get("location")
                dates = _format_date_range(exp.get("startDate"), exp.get("endDate"), exp.get("current"))
                lines.append(f"{title}{f' · {company}' if company else ''}")
                if location or dates:
                    meta = " | ".join([item for item in [location, dates] if item])
                    lines.append(f"  {meta}")
                if exp.get("description"):
                    lines.append(f"  {exp.get('description')}")
                lines.extend([""] * (1 + content_gap_lines))
    elif section == "education":
        education = cv_data.get("education", [])
        if education:
            lines.append(label.upper())
            lines.append("-" * len(label))
            for edu in education:
                degree = edu.get("degree") or "Sin título"
                institution = edu.get("institution")
                location = edu.get("location")
                dates = _format_date_range(edu.get("startDate"), edu.get("endDate"), False)
                lines.append(f"{degree}{f' · {institution}' if institution else ''}")
                if location or dates:
                    meta = " | ".join([item for item in [location, dates] if item])
                    lines.append(f"  {meta}")
                if edu.get("description"):
                    lines.append(f"  {edu.get('description')}")
                lines.extend([""] * (1 + content_gap_lines))
    elif section == "skills":
        skills = [skill.get("name") for skill in cv_data.get("skills", []) if skill.get("name")]
        if skills:
            lines.append(label.upper())
            lines.append("-" * len(label))
            lines.append(" • ".join(skills))
    elif section == "projects":
        projects = cv_data.get("projects", [])
        if projects:
            lines.append(label.upper())
            lines.append("-" * len(label))
            for proj in projects:
                name = proj.get("name") or "Proyecto"
                description = proj.get("description")
                lines.append(name)
                if description:
                    lines.append(f"  {description}")
                if proj.get("url"):
                    lines.append(f"  {proj.get('url')}")
                if proj.get("technologies"):
                    lines.append(f"  Tecnologías: {', '.join(proj.get('technologies'))}")
                lines.extend([""] * (1 + content_gap_lines))
    elif section == "certifications":
        certifications = cv_data.get("certifications", [])
        if certifications:
            lines.append(label.upper())
            lines.append("-" * len(label))
            for cert in certifications:
                title = cert.get("name") or "Certificación"
                issuer = cert.get("issuer")
                date = cert.get("date")
                lines.append(f"{title}{f' · {issuer}' if issuer else ''}{f' ({date})' if date else ''}")
                if cert.get("url"):
                    lines.append(f"  {cert.get('url')}")
            lines.extend([""] * (1 + content_gap_lines))
    elif section == "languages":
        languages = cv_data.get("languages", [])
        if languages:
            lines.append(label.upper())
            lines.append("-" * len(label))
            lines.append(
                " • ".join(
                    [
                        _format_language_entry(lang)
                        for lang in languages
                        if lang.get("language")
                    ]
                )
            )
    elif section == "interests":
        interests = cv_data.get("interests", [])
        if interests:
            lines.append(label.upper())
            lines.append("-" * len(label))
            lines.append(" • ".join([interest.get("name") for interest in interests if interest.get("name")]))
    elif section == "tools":
        tools = cv_data.get("tools", [])
        if tools:
            lines.append(label.upper())
            lines.append("-" * len(label))
            lines.append(" • ".join([tool for tool in tools if tool]))

    return [line for line in lines if line is not None]


def _build_docx_export(cv_data: Dict[str, Any], template_id: str) -> bytes:
    document = Document()
    personal_info = cv_data.get("personalInfo", {})
    layout = _get_layout_config(cv_data)
    colors_config = _resolve_colors(cv_data, template_id)
    fonts = _resolve_docx_fonts(cv_data, template_id)
    section_gap_pt = _px_to_pt(layout["sectionGap"])
    content_gap_pt = _px_to_pt(layout["contentGap"])

    name_heading = document.add_heading(personal_info.get("fullName") or "Sin nombre", level=0)
    _apply_run_style(name_heading.runs[0], fonts["heading"], 20 * layout["fontScale"], colors_config["primary"])
    _apply_paragraph_spacing(name_heading, content_gap_pt)

    if personal_info.get("role"):
        role_paragraph = document.add_paragraph(personal_info.get("role"))
        _apply_run_style(role_paragraph.runs[0], fonts["body"], 11 * layout["fontScale"], colors_config["text"])
        _apply_paragraph_spacing(role_paragraph, content_gap_pt)

    contact_parts = [
        personal_info.get("email"),
        personal_info.get("phone"),
        personal_info.get("location"),
    ]
    contact = " | ".join([item for item in contact_parts if item])
    if contact:
        contact_paragraph = document.add_paragraph(contact)
        _apply_run_style(contact_paragraph.runs[0], fonts["body"], 10 * layout["fontScale"], colors_config["text"])
        _apply_paragraph_spacing(contact_paragraph, content_gap_pt)

    link_parts = []
    if personal_info.get("linkedin"):
        link_parts.append(f"LinkedIn: {personal_info.get('linkedin')}")
    if personal_info.get("github"):
        link_parts.append(f"GitHub: {personal_info.get('github')}")
    if personal_info.get("website"):
        link_parts.append(f"Web: {personal_info.get('website')}")
    if link_parts:
        link_paragraph = document.add_paragraph(" | ".join(link_parts))
        _apply_run_style(link_paragraph.runs[0], fonts["body"], 10 * layout["fontScale"], colors_config["text"])
        _apply_paragraph_spacing(link_paragraph, content_gap_pt)

    for section in TEMPLATE_SECTION_ORDER[template_id]:
        if not _is_section_visible(cv_data, section):
            continue
        _render_section_docx(
            document,
            cv_data,
            section,
            fonts,
            colors_config,
            12 * layout["fontScale"],
            10 * layout["fontScale"],
            section_gap_pt,
            content_gap_pt,
        )

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _render_section_docx(
    document: Document,
    cv_data: Dict[str, Any],
    section: str,
    fonts: Dict[str, str],
    colors_config: Dict[str, str],
    heading_size: float,
    body_size: float,
    section_gap_pt: float,
    content_gap_pt: float,
) -> None:
    label = _get_section_label(cv_data, section)

    if section == "summary":
        summary = cv_data.get("personalInfo", {}).get("summary")
        if summary:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            paragraph = document.add_paragraph(summary)
            _apply_run_style(paragraph.runs[0], fonts["body"], body_size, colors_config["text"])
            _apply_paragraph_spacing(paragraph, content_gap_pt)
    elif section == "experience":
        experiences = cv_data.get("experience", [])
        if experiences:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            for exp in experiences:
                title = exp.get("position") or "Sin título"
                company = exp.get("company")
                dates = _format_date_range(exp.get("startDate"), exp.get("endDate"), exp.get("current"))
                header = f"{title}{f' · {company}' if company else ''}"
                paragraph = document.add_paragraph()
                run = paragraph.add_run(header)
                run.bold = True
                _apply_run_style(run, fonts["heading"], body_size, colors_config["text"])
                _apply_paragraph_spacing(paragraph, content_gap_pt)
                if dates:
                    meta = document.add_paragraph(dates)
                    _apply_run_style(meta.runs[0], fonts["body"], body_size * 0.9, colors_config["text"])
                    _apply_paragraph_spacing(meta, content_gap_pt)
                if exp.get("location"):
                    location = document.add_paragraph(exp.get("location"))
                    _apply_run_style(location.runs[0], fonts["body"], body_size * 0.9, colors_config["text"])
                    _apply_paragraph_spacing(location, content_gap_pt)
                if exp.get("description"):
                    description = document.add_paragraph(exp.get("description"))
                    _apply_run_style(description.runs[0], fonts["body"], body_size, colors_config["text"])
                    _apply_paragraph_spacing(description, content_gap_pt)
    elif section == "education":
        education = cv_data.get("education", [])
        if education:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            for edu in education:
                degree = edu.get("degree") or "Sin título"
                institution = edu.get("institution")
                dates = _format_date_range(edu.get("startDate"), edu.get("endDate"), False)
                header = f"{degree}{f' · {institution}' if institution else ''}"
                paragraph = document.add_paragraph()
                run = paragraph.add_run(header)
                run.bold = True
                _apply_run_style(run, fonts["heading"], body_size, colors_config["text"])
                _apply_paragraph_spacing(paragraph, content_gap_pt)
                if dates:
                    meta = document.add_paragraph(dates)
                    _apply_run_style(meta.runs[0], fonts["body"], body_size * 0.9, colors_config["text"])
                    _apply_paragraph_spacing(meta, content_gap_pt)
                if edu.get("location"):
                    location = document.add_paragraph(edu.get("location"))
                    _apply_run_style(location.runs[0], fonts["body"], body_size * 0.9, colors_config["text"])
                    _apply_paragraph_spacing(location, content_gap_pt)
                if edu.get("description"):
                    description = document.add_paragraph(edu.get("description"))
                    _apply_run_style(description.runs[0], fonts["body"], body_size, colors_config["text"])
                    _apply_paragraph_spacing(description, content_gap_pt)
    elif section == "skills":
        skills = [skill.get("name") for skill in cv_data.get("skills", []) if skill.get("name")]
        if skills:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            paragraph = document.add_paragraph(" • ".join(skills))
            _apply_run_style(paragraph.runs[0], fonts["body"], body_size, colors_config["text"])
            _apply_paragraph_spacing(paragraph, content_gap_pt)
    elif section == "projects":
        projects = cv_data.get("projects", [])
        if projects:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            for proj in projects:
                name = proj.get("name") or "Proyecto"
                paragraph = document.add_paragraph()
                run = paragraph.add_run(name)
                run.bold = True
                _apply_run_style(run, fonts["heading"], body_size, colors_config["text"])
                _apply_paragraph_spacing(paragraph, content_gap_pt)
                if proj.get("description"):
                    description = document.add_paragraph(proj.get("description"))
                    _apply_run_style(description.runs[0], fonts["body"], body_size, colors_config["text"])
                    _apply_paragraph_spacing(description, content_gap_pt)
                if proj.get("url"):
                    url = document.add_paragraph(proj.get("url"))
                    _apply_run_style(url.runs[0], fonts["body"], body_size, colors_config["text"])
                    _apply_paragraph_spacing(url, content_gap_pt)
                if proj.get("technologies"):
                    tech = document.add_paragraph(f"Tecnologías: {', '.join(proj.get('technologies'))}")
                    _apply_run_style(tech.runs[0], fonts["body"], body_size, colors_config["text"])
                    _apply_paragraph_spacing(tech, content_gap_pt)
    elif section == "certifications":
        certifications = cv_data.get("certifications", [])
        if certifications:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            for cert in certifications:
                title = cert.get("name") or "Certificación"
                issuer = cert.get("issuer")
                date = cert.get("date")
                header = f"{title}{f' · {issuer}' if issuer else ''}{f' ({date})' if date else ''}"
                paragraph = document.add_paragraph(header)
                _apply_run_style(paragraph.runs[0], fonts["body"], body_size, colors_config["text"])
                _apply_paragraph_spacing(paragraph, content_gap_pt)
                if cert.get("url"):
                    url = document.add_paragraph(cert.get("url"))
                    _apply_run_style(url.runs[0], fonts["body"], body_size, colors_config["text"])
                    _apply_paragraph_spacing(url, content_gap_pt)
    elif section == "languages":
        languages = cv_data.get("languages", [])
        if languages:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            content = " • ".join(
                [
                    _format_language_entry(lang)
                    for lang in languages
                    if lang.get("language")
                ]
            )
            paragraph = document.add_paragraph(content)
            _apply_run_style(paragraph.runs[0], fonts["body"], body_size, colors_config["text"])
            _apply_paragraph_spacing(paragraph, content_gap_pt)
    elif section == "interests":
        interests = cv_data.get("interests", [])
        if interests:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            content = " • ".join([interest.get("name") for interest in interests if interest.get("name")])
            paragraph = document.add_paragraph(content)
            _apply_run_style(paragraph.runs[0], fonts["body"], body_size, colors_config["text"])
            _apply_paragraph_spacing(paragraph, content_gap_pt)
    elif section == "tools":
        tools = cv_data.get("tools", [])
        if tools:
            heading = document.add_heading(label, level=1)
            _apply_run_style(heading.runs[0], fonts["heading"], heading_size, colors_config["accent"])
            _apply_paragraph_spacing(heading, section_gap_pt)
            content = " • ".join([tool for tool in tools if tool])
            paragraph = document.add_paragraph(content)
            _apply_run_style(paragraph.runs[0], fonts["body"], body_size, colors_config["text"])
            _apply_paragraph_spacing(paragraph, content_gap_pt)


def _build_pdf_export(cv_data: Dict[str, Any], template_id: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    layout = _get_layout_config(cv_data)
    colors_config = _resolve_colors(cv_data, template_id)
    pdf_fonts = _resolve_pdf_fonts(cv_data, template_id)
    style = {
        "heading_font": pdf_fonts["heading"],
        "body_font": pdf_fonts["body"],
        "accent": colors.HexColor(colors_config["accent"]),
        "text": colors.HexColor(colors_config["text"]),
        "section_gap": _px_to_pt(layout["sectionGap"]),
        "content_gap": _px_to_pt(layout["contentGap"]),
        "font_scale": layout["fontScale"],
    }

    pdf.setTitle("CV Export")
    cursor_y = height - 72

    personal_info = cv_data.get("personalInfo", {})
    full_name = personal_info.get("fullName") or "Sin nombre"
    role = personal_info.get("role")

    cursor_y = _draw_wrapped_text(
        pdf,
        full_name,
        cursor_y,
        style["heading_font"],
        int(18 * style["font_scale"]),
        style["accent"],
    )

    if role:
        cursor_y = _draw_wrapped_text(pdf, role, cursor_y, style["body_font"], int(12 * style["font_scale"]), style["text"])

    contact_parts = [
        personal_info.get("email"),
        personal_info.get("phone"),
        personal_info.get("location"),
    ]
    contact = " | ".join([item for item in contact_parts if item])
    if contact:
        cursor_y = _draw_wrapped_text(pdf, contact, cursor_y, style["body_font"], int(10 * style["font_scale"]), style["text"])

    link_parts = []
    if personal_info.get("linkedin"):
        link_parts.append(f"LinkedIn: {personal_info.get('linkedin')}")
    if personal_info.get("github"):
        link_parts.append(f"GitHub: {personal_info.get('github')}")
    if personal_info.get("website"):
        link_parts.append(f"Web: {personal_info.get('website')}")
    if link_parts:
        cursor_y = _draw_wrapped_text(pdf, " | ".join(link_parts), cursor_y, style["body_font"], int(10 * style["font_scale"]), style["text"])

    cursor_y -= style["section_gap"]

    for section in TEMPLATE_SECTION_ORDER[template_id]:
        if not _is_section_visible(cv_data, section):
            continue
        cursor_y = _render_section_pdf(pdf, cv_data, section, style, cursor_y)

    pdf.showPage()
    pdf.save()
    buffer.seek(0)
    return buffer.read()


def _render_section_pdf(
    pdf: canvas.Canvas,
    cv_data: Dict[str, Any],
    section: str,
    style: Dict[str, Any],
    cursor_y: float,
) -> float:
    label = _get_section_label(cv_data, section)
    item_gap = max(4, style["content_gap"] * 0.5)
    cursor_y = _draw_wrapped_text(pdf, label, cursor_y, style["heading_font"], int(12 * style["font_scale"]), style["accent"])

    if section == "summary":
        summary = cv_data.get("personalInfo", {}).get("summary")
        if summary:
            cursor_y = _draw_wrapped_text(pdf, summary, cursor_y, style["body_font"], int(10 * style["font_scale"]), style["text"])
    elif section == "experience":
        experiences = cv_data.get("experience", [])
        for exp in experiences:
            title = exp.get("position") or "Sin título"
            company = exp.get("company")
            header = f"{title}{f' · {company}' if company else ''}"
            cursor_y = _draw_wrapped_text(pdf, header, cursor_y, style["heading_font"], int(10 * style["font_scale"]), style["text"])
            dates = _format_date_range(exp.get("startDate"), exp.get("endDate"), exp.get("current"))
            meta = " | ".join([item for item in [exp.get("location"), dates] if item])
            if meta:
                cursor_y = _draw_wrapped_text(pdf, meta, cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            if exp.get("description"):
                cursor_y = _draw_wrapped_text(pdf, exp.get("description"), cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            cursor_y -= item_gap
    elif section == "education":
        education = cv_data.get("education", [])
        for edu in education:
            degree = edu.get("degree") or "Sin título"
            institution = edu.get("institution")
            header = f"{degree}{f' · {institution}' if institution else ''}"
            cursor_y = _draw_wrapped_text(pdf, header, cursor_y, style["heading_font"], int(10 * style["font_scale"]), style["text"])
            dates = _format_date_range(edu.get("startDate"), edu.get("endDate"), False)
            meta = " | ".join([item for item in [edu.get("location"), dates] if item])
            if meta:
                cursor_y = _draw_wrapped_text(pdf, meta, cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            if edu.get("description"):
                cursor_y = _draw_wrapped_text(pdf, edu.get("description"), cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            cursor_y -= item_gap
    elif section == "skills":
        skills = [skill.get("name") for skill in cv_data.get("skills", []) if skill.get("name")]
        if skills:
            cursor_y = _draw_wrapped_text(pdf, " • ".join(skills), cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
    elif section == "projects":
        projects = cv_data.get("projects", [])
        for proj in projects:
            name = proj.get("name") or "Proyecto"
            cursor_y = _draw_wrapped_text(pdf, name, cursor_y, style["heading_font"], int(10 * style["font_scale"]), style["text"])
            if proj.get("description"):
                cursor_y = _draw_wrapped_text(pdf, proj.get("description"), cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            if proj.get("url"):
                cursor_y = _draw_wrapped_text(pdf, proj.get("url"), cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            if proj.get("technologies"):
                tech = f"Tecnologías: {', '.join(proj.get('technologies'))}"
                cursor_y = _draw_wrapped_text(pdf, tech, cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            cursor_y -= item_gap
    elif section == "certifications":
        certifications = cv_data.get("certifications", [])
        for cert in certifications:
            title = cert.get("name") or "Certificación"
            issuer = cert.get("issuer")
            date = cert.get("date")
            header = f"{title}{f' · {issuer}' if issuer else ''}{f' ({date})' if date else ''}"
            cursor_y = _draw_wrapped_text(pdf, header, cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            if cert.get("url"):
                cursor_y = _draw_wrapped_text(pdf, cert.get("url"), cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
            cursor_y -= item_gap
    elif section == "languages":
        languages = cv_data.get("languages", [])
        if languages:
            content = " • ".join([_format_language_entry(lang) for lang in languages if lang.get("language")])
            cursor_y = _draw_wrapped_text(pdf, content, cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
    elif section == "interests":
        interests = cv_data.get("interests", [])
        if interests:
            content = " • ".join([interest.get("name") for interest in interests if interest.get("name")])
            cursor_y = _draw_wrapped_text(pdf, content, cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])
    elif section == "tools":
        tools = cv_data.get("tools", [])
        if tools:
            content = " • ".join([tool for tool in tools if tool])
            cursor_y = _draw_wrapped_text(pdf, content, cursor_y, style["body_font"], int(9 * style["font_scale"]), style["text"])

    return cursor_y - style["section_gap"]


def _draw_wrapped_text(
    pdf: canvas.Canvas,
    text: str,
    cursor_y: float,
    font_name: str,
    font_size: int,
    color: colors.Color,
) -> float:
    if cursor_y < 72:
        pdf.showPage()
        cursor_y = letter[1] - 72

    pdf.setFont(font_name, font_size)
    pdf.setFillColor(color)

    max_width = letter[0] - 144
    wrapped = _wrap_text(text, font_name, font_size, max_width)
    for line in wrapped:
        if cursor_y < 72:
            pdf.showPage()
            cursor_y = letter[1] - 72
            pdf.setFont(font_name, font_size)
            pdf.setFillColor(color)
        pdf.drawString(72, cursor_y, line)
        cursor_y -= font_size + 2
    return cursor_y


def _wrap_text(text: str, font_name: str, font_size: int, max_width: float) -> List[str]:
    if not text:
        return []
    approx_char_width = font_size * 0.5
    max_chars = max(int(max_width / approx_char_width), 10)
    return textwrap.wrap(text, width=max_chars)


def _format_language_entry(language: Dict[str, Any]) -> str:
    fluency = language.get("fluency")
    if fluency:
        return f"{language.get('language')} ({fluency})"
    return str(language.get("language"))


def _format_date_range(start: str, end: str, current: bool) -> str:
    if not start and not end and not current:
        return ""
    if current:
        return f"{start or ''} - Presente".strip(" -")
    if start and end:
        return f"{start} - {end}"
    return start or end or ""
